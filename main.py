from flask import Flask, render_template_string, request
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import base64
import os
import uuid
from datetime import datetime
from dateutil.relativedelta import relativedelta
import requests
import io
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# --- Configuration ---
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
TELEGRAM_CHAT_ID = os.getenv("TELEGRAM_CHAT_ID")

app = Flask(__name__)

# --- Helper Function for Date Formatting ---
def format_date_with_suffix(d):
    """ Formats a date object into '03rd day of August 2025' style """
    day = d.day
    if 4 <= day <= 20 or 24 <= day <= 30:
        suffix = "th"
    else:
        suffix = ["st", "nd", "rd"][day % 10 - 1]
    return f"{day:02d}{suffix} day of {d.strftime('%B %Y')}"

# --- Telegram Bot Function ---
def send_file_to_telegram(document_stream, filename, caption):
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        return False, "Credentials missing"

    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument"
    document_stream.seek(0)
    
    try:
        files = {
            'document': (filename, document_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        data = {
            'chat_id': TELEGRAM_CHAT_ID,
            'caption': caption
        }
        
        response = requests.post(url, data=data, files=files)
        response_json = response.json()
        
        if response_json.get("ok"):
            return True, "Sent successfully"
        else:
            return False, response_json.get("description", "Unknown error")
            
    except Exception as e:
        return False, str(e)

# --- Word Document Generation Logic ---
def create_word_agreement(client_data):
    doc = Document()

    # --- PAGE 1 SETUP ---
    section_page1 = doc.sections[0]
    section_page1.page_height = Inches(14.0)
    section_page1.page_width = Inches(8.5)
    section_page1.left_margin = Cm(3.0)
    section_page1.right_margin = Cm(1.5)
    
    def add_paragraph_with_runs(texts_and_formats, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=16):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = alignment
        for text, is_bold in texts_and_formats:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.bold = is_bold
        return p

    def add_formatted_paragraph(text, size=16, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = align
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        return p

    # --- Prepare Data ---
    start_date = datetime.strptime(client_data['start_date'], '%Y-%m-%d').date()
    stay_months = int(client_data.get('stay_months') or 0)
    
    next_month_date = start_date + relativedelta(months=+stay_months)
    first_day_of_next_month = next_month_date.replace(day=1)
    end_date = first_day_of_next_month - relativedelta(days=1)
    
    start_date_str = format_date_with_suffix(start_date)
    end_date_str = format_date_with_suffix(end_date)
    
    full_name = f"{client_data['first_name']} {client_data['last_name']}"
    
    # Format Addresses
    full_address = f"{client_data['address']}, {client_data['permanent_district']}, {client_data['permanent_state']} - {client_data['permanent_pincode']}"
    
    # Since all office fields are now mandatory:
    full_office_address = f"{client_data['office_address']}, {client_data['office_district']}, {client_data['office_state']} - {client_data['office_pincode']}"
    
    rent_in_words = f"Rupees {num2words(int(client_data['rent_price']), lang='en_IN').title()} Only"
    deposit_in_words = f"Rupees {num2words(int(client_data['security_deposit']), lang='en_IN').title()} Only"

    # --- PAGE 1 CONTENT ---
    for _ in range(17): doc.add_paragraph()

    add_formatted_paragraph('PAYING GUEST AGREEMENT', size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_paragraph_with_runs([
        ("THIS AGREEMENT is made and entered in to at Mumbai this ", False),
        (f"{start_date_str} BETWEEN: MR. JASMEET SINGH", True),
        (", residing at ", False),
        ("303/B wing, Palatial Heights, Chandivali Farm Rd, Chandivali, Powai, Mumbai, Maharashtra 400072", True),
        (", Hereinafter referred to as ", False),
        ("“CARETAKER”", True),
        (" (which expression shall mean and include his heirs, executors, administrators and assigns) of the ", False),
        ("ONE PART", True)
    ], font_size=16, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    doc.add_page_break()

    # --- PAGE 2 & 3 ---
    legal_section = doc.sections[-1]
    legal_section.page_height = Inches(14.0)
    legal_section.page_width = Inches(8.5)
    legal_section.left_margin = Cm(3.0)
    legal_section.right_margin = Cm(1.5)

    add_formatted_paragraph('AND', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    font_size_main = Pt(14)
    p_details = doc.add_paragraph()
    p_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_details.paragraph_format.space_before = Pt(12)
    p_details.paragraph_format.space_after = Pt(0)

    def add_run_to_details(text, bold=False):
        run = p_details.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = font_size_main
        run.bold = bold

    add_run_to_details(f"{client_data['salutation']}. ", bold=True)
    add_run_to_details(full_name, bold=True)
    add_run_to_details(f", aged {client_data['age']} years, an adult, ")
    add_run_to_details("Indian Inhabitant permanently residing at: ")
    add_run_to_details(full_address, bold=True)
    add_run_to_details(" Having Aadhar card No. ")
    add_run_to_details(client_data['aadhar_no'], bold=True)
    add_run_to_details("\n")

    add_run_to_details("Emergency Contact:\n")
    add_run_to_details("(1) ")
    add_run_to_details(client_data['ref1_name'], bold=True)
    add_run_to_details(" Ph- ")
    add_run_to_details(client_data['ref1_number'], bold=True)
    add_run_to_details("\n")
    add_run_to_details("(2) ")
    add_run_to_details(client_data['ref2_name'], bold=True)
    add_run_to_details(" Ph- ")
    add_run_to_details(client_data['ref2_number'], bold=True)
    add_run_to_details("\n")
    
    add_run_to_details("Office Address: ")
    add_run_to_details(full_office_address, bold=True)
    add_run_to_details("\n")
        
    add_run_to_details("Email ID: ")
    add_run_to_details(client_data['email_id'], bold=True)
    add_run_to_details("\n")
    
    p_last = doc.add_paragraph()
    p_last.paragraph_format.space_before = Pt(0)
    p_last.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_last.add_run("Hereinafter referred to as the ").font.size = font_size_main
    run = p_last.add_run("“PAYING GUEST” "); run.bold = True; run.font.size = font_size_main
    p_last.add_run("(which expression shall mean and include his heirs, executors, administrators and assigns) of the ").font.size = font_size_main
    run = p_last.add_run("SECOND PART."); run.bold = True; run.font.size = font_size_main
    for run in p_last.runs: run.font.name = 'Times New Roman'

    add_paragraph_with_runs([("WHEREAS", True), (" the party of the one Part is the Host in respect of premises situate at ", False), (client_data['rented_address'], True), (", hereinafter for the sake of brevity referred to as the “Said Room Premises”.", False)], font_size=14)
    add_paragraph_with_runs([("AND WHEREAS", True), (" the Paying Guests are in need of temporary furnished accommodation and has approached and requested to the owner to permit the said Paying Guest the use of the “Said Room Premises” together with the fixtures, fittings, furniture’s and amenities for residential purposes for a temporary period. AND WHEREAS, the Host has agreed on certain terms and conditions which the parties have mutually agreed themselves as under.", False)], font_size=14)
    doc.add_paragraph()

    clauses = [
        [("The Host has permitted the Paying Guest the Use of part bathrooms in the “Said room Premises” situated at ", False), (client_data["rented_address"], True), (" together with fixtures, fittings, furniture and amenities for the purpose of providing temporary residential accommodation on paying guest basis.", False)],
        [("This Agreement shall be on monthly basis commencing from ", False), (start_date_str, True), (" to ", False), (end_date_str, True)],
        [('The Paying Guest shall pay the monthly rent between the 1st and 5th day of every month. Any delay beyond the 5th day shall attract a late payment charge of ₹200 (Rupees Two Hundred) per day until the rent is cleared. Upon vacating the “Said Room Premises,” a sum of ₹500 (Rupees Five Hundred) shall be deducted from the Security Deposit towards room cleaning charges, and the remaining balance of the deposit, if any, shall be refunded after adjustment of all dues or damages, if applicable.', True)],
        [(f"That the Paying Guest shall pay ", False), (f"Rs. {int(client_data['security_deposit']):}/- ({deposit_in_words})", True), (" as a refundable security deposit amount to the Caretaker. which will be returned to the Paying Guest on vacating the “ Said Room Premises” for which ", False), ("ONE MONTH", True), (" notice is required.", False)],
        [(f"That the Paying Guest shall pay to the caretaker of ", False), (f"Rs. {int(client_data['rent_price']):}/- ({rent_in_words})", True), (" towards the compensation charges for the use of the “Said Room Premises” together with the use of the fixtures, fittings, furniture and amenities and which is not including Electricity Charges (actual) to be shared by all PG’s as also maid charges.", False)],
        [("The Paying Guest shall keep the “Said Room Premises” in good condition and comply with all the rules and regulations required in this regard.", False)],
        [("The paying Guest shall not carry out any addition or alterations in the “Said Room Premises”.", False)],
        [("The “Said Room Premises” shall be used by the Paying Guest Only for lawful purpose of residential stay. The said premises shall not be used for any other purpose/s by the Paying Guest. The Caretaker shall restrain the access to the “Said Room Premises” if the paying guest misuses the premises or commits any illegal act or criminal act or disturbs the neighbors or the society.", False)],
        [("The Paying Guest hereby covenants and agrees that they shall not use the address of the “Said Room Premises” for obtaining, applying for, or registering any government-issued identification, documentation, or services, including but not limited to: Ration Card, Gas Connection, Aadhaar Card, PAN Card, Voter ID Card, Driving License, Bank Loan or Online Loan documentation, Any other government-recognized proof of residence.", False)],
        [("The paying guest shall not bring any visitors to the premises except with the permission of the Caretaker.", False)],
        [("The Caretaker of his representatives shall have the lock and key of the “Said Room Premises” and have the right to enter the said room for the purpose of inspection or any other purpose/s at all reasonable hours.", False)],
        [("The Notice period for termination of this paying guest by either party is ONE MONTH. (The paying Guest have no right to vacate the said premises before 3 months from the commencing of this agreement) .(i.e. 3 months locking period)", False)],
        [("This agreement does not bestow any right, title, possession or interest of whatsoever nature in the Room / Flat to the Paying Guest.", False)]
    ]

    for i, clause_parts in enumerate(clauses):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        if i == 3: p.paragraph_format.page_break_before = True
        for text, is_bold in clause_parts:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = is_bold
        if i != 12: doc.add_paragraph()

    signature_page_section = doc.sections[-1]
    signature_page_section.left_margin = Cm(3.0)
    signature_page_section.right_margin = Cm(1.5)

    add_formatted_paragraph("IN WITNESS WHEREOF the parties have hereto hereinto set their respective hands on the day and year first hereinabove mentioned.", size=14)
    for _ in range(3): doc.add_paragraph()

    add_paragraph_with_runs([('SIGNED AND DELIVERED for\nThe Caretaker by withinnamed\n', False), ('Mr. Jasmeet Singh', True)], alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=14)
    doc.add_paragraph()
    add_formatted_paragraph('In the presence of ………………….', size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    for _ in range(5): doc.add_paragraph()
    
    add_paragraph_with_runs([(f'SIGNED AND DELIVERED for\nThe paying Guest by withinnamed\n', False), (f'{client_data["salutation"]}. ', True), (full_name, True)], alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=14)
    
    sig_paragraph = doc.add_paragraph()
    sig_run = sig_paragraph.add_run()
    if os.path.exists(client_data['signature_path']):
        sig_run.add_picture(client_data['signature_path'], width=Inches(2.0))
    sig_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_formatted_paragraph('In the presence of ………………….', size=14, align=WD_ALIGN_PARAGRAPH.LEFT)

    document_stream = io.BytesIO()
    doc.save(document_stream)
    return document_stream

# --- HTML Template ---
# Updated: strict 'required' attributes on ALL fields including Office and Email
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Paying Guest Agreement Form</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body { font-family: 'Inter', sans-serif; }
        .signature-pad { border: 2px dashed #ccc; border-radius: 8px; cursor: crosshair; }
        .section-title { font-size: 1.125rem; font-weight: 600; color: #1f2937; border-bottom: 1px solid #e5e7eb; padding-bottom: 0.5rem; margin-bottom: 1rem; }
    </style>
</head>
<body class="bg-gray-100 flex items-center justify-center min-h-screen py-8">
    <div id="form-container" class="w-full max-w-2xl p-8 space-y-6 bg-white rounded-lg shadow-md m-4">
        <h1 class="text-3xl font-bold text-center text-gray-800">Paying Guest Details Form</h1>
        <p class="text-center text-gray-600">Please fill in <strong>ALL</strong> details below. No field can be left blank.</p>
        
        <form action="/submit" method="post" id="agreementForm" class="space-y-8">
            
            <div>
                <h2 class="section-title">Personal Details</h2>
                <div class="space-y-6 mt-4">
                    
                    <div>
                        <label for="salutation" class="block text-sm font-medium text-gray-700">Salutation *</label>
                        <select id="salutation" name="salutation" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                            <option value="Ms">Ms.</option>
                            <option value="Mr">Mr.</option>
                        </select>
                    </div>

                    <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                        <div>
                            <label for="first_name" class="block text-sm font-medium text-gray-700">First Name *</label>
                            <input type="text" id="first_name" name="first_name" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                        <div>
                            <label for="last_name" class="block text-sm font-medium text-gray-700">Last Name *</label>
                            <input type="text" id="last_name" name="last_name" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                    </div>
                    <div>
                        <label for="age" class="block text-sm font-medium text-gray-700">Age *</label>
                        <input type="number" id="age" name="age" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>
                    
                    <div>
                        <label for="address" class="block text-sm font-medium text-gray-700">Permanent Address (Street/Building) *</label>
                        <input type="text" id="address" name="address" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>

                    <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                         <div>
                            <label for="permanent_district" class="block text-sm font-medium text-gray-700">District *</label>
                            <input type="text" id="permanent_district" name="permanent_district" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                        <div>
                            <label for="permanent_state" class="block text-sm font-medium text-gray-700">State *</label>
                            <select id="permanent_state" name="permanent_state" required class="state-dropdown mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                                <option value="">Select State</option>
                            </select>
                        </div>
                    </div>

                    <div>
                        <label for="permanent_pincode" class="block text-sm font-medium text-gray-700">Pincode (Permanent) *</label>
                        <input type="text" id="permanent_pincode" name="permanent_pincode" required pattern="[0-9]{6}" title="Enter a 6-digit pincode" class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>
                    
                    <div>
                        <label for="aadhar_no" class="block text-sm font-medium text-gray-700">Aadhar Card Number *</label>
                        <input type="text" id="aadhar_no" name="aadhar_no" required pattern="[0-9]{12}" title="Enter a 12-digit Aadhar number" class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>

                    <div>
                        <label for="office_address" class="block text-sm font-medium text-gray-700">Office Address (Enter 'N/A' if none) *</label>
                        <input type="text" id="office_address" name="office_address" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>

                    <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                         <div>
                            <label for="office_district" class="block text-sm font-medium text-gray-700">District (Office) *</label>
                            <input type="text" id="office_district" name="office_district" required placeholder="or N/A" class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                        <div>
                            <label for="office_state" class="block text-sm font-medium text-gray-700">State (Office) *</label>
                            <select id="office_state" name="office_state" required class="state-dropdown mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                                <option value="">Select State</option>
                            </select>
                        </div>
                    </div>

                    <div>
                        <label for="office_pincode" class="block text-sm font-medium text-gray-700">Pincode (Office) *</label>
                        <input type="text" id="office_pincode" name="office_pincode" required pattern="[0-9]{6}" title="Enter 6-digit pincode (or 000000 if N/A)" class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>

                    <div>
                        <label for="email_id" class="block text-sm font-medium text-gray-700">Email ID *</label>
                        <input type="email" id="email_id" name="email_id" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>
                </div>
            </div>

            <div>
                <h2 class="section-title">Reference Contacts</h2>
                <div class="space-y-6 mt-4">
                    <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                        <div>
                            <label for="ref1_name" class="block text-sm font-medium text-gray-700">Reference 1 Name *</label>
                            <input type="text" id="ref1_name" name="ref1_name" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                        <div>
                            <label for="ref1_number" class="block text-sm font-medium text-gray-700">Reference 1 Number *</label>
                            <input type="tel" id="ref1_number" name="ref1_number" required pattern="[0-9]{10}" title="Enter a 10-digit mobile number" class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                    </div>
                    <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                        <div>
                            <label for="ref2_name" class="block text-sm font-medium text-gray-700">Reference 2 Name *</label>
                            <input type="text" id="ref2_name" name="ref2_name" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                        <div>
                            <label for="ref2_number" class="block text-sm font-medium text-gray-700">Reference 2 Number *</label>
                            <input type="tel" id="ref2_number" name="ref2_number" required pattern="[0-9]{10}" title="Enter a 10-digit mobile number" class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                    </div>
                </div>
            </div>

            <div>
                <h2 class="section-title">Agreement Terms</h2>
                <div class="space-y-6 mt-4">
                    <div>
                        <label for="rented_address" class="block text-sm font-medium text-gray-700">Rented Property Address *</label>
                        <input type="text" id="rented_address" name="rented_address" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>
                    <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                        <div>
                            <label for="rent_price" class="block text-sm font-medium text-gray-700">Monthly Rent (INR) *</label>
                            <input type="number" id="rent_price" name="rent_price" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                        <div>
                            <label for="security_deposit" class="block text-sm font-medium text-gray-700">Security Deposit (INR) *</label>
                            <input type="number" id="security_deposit" name="security_deposit" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                        </div>
                    </div>
                    <div>
                        <label for="start_date" class="block text-sm font-medium text-gray-700">Agreement Start Date *</label>
                        <input type="date" id="start_date" name="start_date" required class="mt-1 block w-full px-3 py-2 bg-white border border-gray-300 rounded-md shadow-sm">
                    </div>
                </div>
            </div>

            <div>
                <h2 class="section-title">Signature</h2>
                <div class="mt-1 relative">
                    <canvas id="signature-pad" class="signature-pad w-full h-48 bg-gray-50"></canvas>
                    <button type="button" id="clear-signature" class="absolute top-2 right-2 px-3 py-1 text-sm text-white bg-red-600 rounded-md hover:bg-red-700">Clear</button>
                </div>
                <input type="hidden" name="signature" id="signature-data" required>
            </div>

            <button type="submit" class="w-full flex justify-center py-3 px-4 border border-transparent rounded-md shadow-sm text-sm font-medium text-white bg-indigo-600 hover:bg-indigo-700">
                Submit Agreement
            </button>
        </form>
    </div>
    
    <script src="https://cdn.jsdelivr.net/npm/signature_pad@4.0.0/dist/signature_pad.umd.min.js"></script>
    <script>
        const indianStates = [
            "Andhra Pradesh", "Arunachal Pradesh", "Assam", "Bihar", "Chhattisgarh", "Goa", "Gujarat", 
            "Haryana", "Himachal Pradesh", "Jharkhand", "Karnataka", "Kerala", "Madhya Pradesh", 
            "Maharashtra", "Manipur", "Meghalaya", "Mizoram", "Nagaland", "Odisha", "Punjab", 
            "Rajasthan", "Sikkim", "Tamil Nadu", "Telangana", "Tripura", "Uttar Pradesh", 
            "Uttarakhand", "West Bengal", "Andaman and Nicobar Islands", "Chandigarh", 
            "Dadra and Nagar Haveli and Daman and Diu", "Delhi", "Jammu and Kashmir", "Ladakh", 
            "Lakshadweep", "Puducherry", "N/A"
        ];

        document.querySelectorAll('.state-dropdown').forEach(dropdown => {
            indianStates.forEach(state => {
                const option = document.createElement('option');
                option.value = state;
                option.textContent = state;
                dropdown.appendChild(option);
            });
        });

        const canvas = document.getElementById('signature-pad');
        const signaturePad = new SignaturePad(canvas, { backgroundColor: 'rgb(249, 250, 251)' });

        document.getElementById('clear-signature').addEventListener('click', function () {
            signaturePad.clear();
        });

        document.getElementById('agreementForm').addEventListener('submit', function (event) {
            if (signaturePad.isEmpty()) {
                alert("Please provide a signature.");
                event.preventDefault();
                return;
            }
            document.getElementById('signature-data').value = signaturePad.toDataURL('image/png');
        });

        const today = new Date();
        const yyyy = today.getFullYear();
        const mm = String(today.getMonth() + 1).padStart(2, '0');
        const dd = String(today.getDate()).padStart(2, '0');
        document.getElementById('start_date').value = `${yyyy}-${mm}-${dd}`;
    </script>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/submit', methods=['POST'])
def submit():
    signature_path = None
    try:
        # STRICT SERVER-SIDE VALIDATION: Check if any expected field is missing or empty
        required_fields = [
            'salutation', 'first_name', 'last_name', 'age', 
            'address', 'permanent_district', 'permanent_state', 'permanent_pincode',
            'aadhar_no', 
            'office_address', 'office_district', 'office_state', 'office_pincode', 
            'email_id',
            'ref1_name', 'ref1_number', 'ref2_name', 'ref2_number',
            'rented_address', 'rent_price', 'security_deposit', 'start_date', 
            'signature'
        ]

        # Check every field exists and is not just whitespace
        for field in required_fields:
            if not request.form.get(field) or not request.form.get(field).strip():
                return f"""
                <div style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
                    <h1 style="color: #dc3545;">Submission Failed</h1>
                    <p>The field <strong>{field.replace('_', ' ').title()}</strong> cannot be empty.</p>
                    <a href="/">Go Back</a>
                </div>
                """, 400

        # If validation passes, collect data
        client_data = {
            'salutation': request.form['salutation'],
            'first_name': request.form['first_name'],
            'last_name': request.form['last_name'],
            'age': request.form['age'],
            'address': request.form['address'],
            'permanent_district': request.form['permanent_district'],
            'permanent_state': request.form['permanent_state'],
            'permanent_pincode': request.form['permanent_pincode'],
            'aadhar_no': request.form['aadhar_no'],
            'office_address': request.form['office_address'],
            'office_district': request.form['office_district'],
            'office_state': request.form['office_state'],
            'office_pincode': request.form['office_pincode'],
            'email_id': request.form['email_id'],
            'ref1_name': request.form['ref1_name'],
            'ref1_number': request.form['ref1_number'],
            'ref2_name': request.form['ref2_name'],
            'ref2_number': request.form['ref2_number'],
            'rented_address': request.form['rented_address'],
            'rent_price': request.form['rent_price'],
            'security_deposit': request.form['security_deposit'],
            'start_date': request.form['start_date'],
            'stay_months': '11',
            'signature_data_url': request.form['signature']
        }

        # --- Process Signature ---
        header, encoded = client_data['signature_data_url'].split(",", 1)
        signature_data = base64.b64decode(encoded)
        signature_filename = f"signature_{uuid.uuid4().hex}.png"
        signature_path = os.path.join('/tmp', signature_filename)
        with open(signature_path, "wb") as f:
            f.write(signature_data)
        client_data['signature_path'] = signature_path
        
        # --- Generate and Send ---
        document_stream = create_word_agreement(client_data)
        full_name = f"{client_data['first_name']} {client_data['last_name']}"
        filename = f"Agreement_{full_name.replace(' ', '_')}.docx"
        caption = f"New agreement submitted by: {client_data['salutation']}. {full_name}\nAadhar: {client_data['aadhar_no']}"
        
        success, message = send_file_to_telegram(document_stream, filename, caption)

        if success:
            return """
                <div style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
                    <h1 style="color: #28a745;">Agreement Sent!</h1>
                    <p style="font-size: 1.2em;">The agreement has been generated and sent to the administrator via Telegram.</p>
                    <a href="/">Go Back</a>
                </div>
            """
        else:
            return f"""
                <div style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
                    <h1 style="color: #dc3545;">Submission Failed</h1>
                    <p>We could not send the document to Telegram.</p>
                    <p style="background: #eee; padding: 10px; display: inline-block;">Error: {message}</p>
                    <br><br>
                    <a href="/">Try Again</a>
                </div>
            """, 500

    except Exception as e:
        print(f"Error in submit route: {e}")
        return f"An error occurred: {e}", 500
    finally:
        if signature_path and os.path.exists(signature_path):
            os.remove(signature_path)

if __name__ == '__main__':
    app.run(debug=True)